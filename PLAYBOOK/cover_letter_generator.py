#!/usr/bin/env python3

import argparse
import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_calibri_font(run, size_pt, bold=False, italic=False):
    run.font.name = 'Calibri'
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic

    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rFonts.set(qn('w:cs'), 'Calibri')


def clean_markdown_text(text):
    """Strip simple markdown formatting used in cover letter files."""
    text = re.sub(r'^\s*#+\s+', '', text)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    return text.strip()


def extract_body_paragraphs(content):
    """Return only real cover-letter paragraphs, skipping markdown scaffolding."""
    paragraphs = []

    for raw_paragraph in content.split('\n\n'):
        paragraph = raw_paragraph.strip()
        if not paragraph:
            continue
        if re.fullmatch(r'-{3,}', paragraph):
            continue
        if paragraph.startswith('#'):
            continue

        cleaned = clean_markdown_text(paragraph)
        if cleaned:
            paragraphs.append(cleaned)

    return paragraphs


def create_cover_letter_from_markdown(md_file_path, output_path):
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    paragraphs = extract_body_paragraphs(content)

    for para_text in paragraphs:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_after = Pt(12)
        para.paragraph_format.line_spacing = 1.15
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

        run = para.add_run(para_text)
        set_calibri_font(run, 11)

    doc.save(output_path)
    print(f"✅ Cover letter generated: {output_path}")
    return True


def main():
    parser = argparse.ArgumentParser(
        description='Generate cover letter from markdown',
        epilog='Example: python3 cover_letter_generator.py -i COVERLETTER.md -o Coverletter.docx'
    )
    parser.add_argument('--input', '-i', required=True, help='Input COVERLETTER.md file')
    parser.add_argument('--output', '-o', required=True, help='Output Coverletter.docx file')

    args = parser.parse_args()

    try:
        create_cover_letter_from_markdown(args.input, args.output)
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
        return 1

    return 0


if __name__ == '__main__':
    exit(main())
